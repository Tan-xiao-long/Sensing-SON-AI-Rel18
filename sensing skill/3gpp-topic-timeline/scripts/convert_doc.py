#!/usr/bin/env python3
"""Convert .doc/.docx to plain text, portably.

Strategy (first available wins):
  .docx -> python-docx if installed, else stdlib zipfile+XML strip
  .doc  -> libreoffice --headless, else antiword, else built-in binary extraction

Usage:  python convert_doc.py INPUT [-o OUTPUT.txt]
Import: from convert_doc import to_text
"""
import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile


def _docx_via_zip(path: str) -> str:
    """Extract text from .docx using only the stdlib (no python-docx)."""
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8", "replace")
    # paragraph and cell boundaries -> newlines, tabs for cells
    xml = re.sub(r"</w:p>", "\n", xml)
    xml = re.sub(r"</w:tc>", "\t", xml)
    xml = re.sub(r"<w:tab[^>]*/>", "\t", xml)
    xml = re.sub(r"<w:br[^>]*/>", "\n", xml)
    text = re.sub(r"<[^>]+>", "", xml)
    for a, b in (("&amp;", "&"), ("&lt;", "<"), ("&gt;", ">"), ("&quot;", '"'), ("&apos;", "'")):
        text = text.replace(a, b)
    return text


def _docx_text(path: str) -> str:
    try:
        import docx  # type: ignore
        d = docx.Document(path)
        parts = []
        for p in d.paragraphs:
            parts.append(p.text)
        for t in d.tables:
            for row in t.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception:
        return _docx_via_zip(path)


def _doc_via_tool(path: str, outdir: str):
    """Try libreoffice / antiword for legacy .doc. Returns text or None."""
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if soffice:
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", outdir, path],
                check=True, capture_output=True, timeout=300,
            )
            cand = os.path.join(outdir, os.path.splitext(os.path.basename(path))[0] + ".txt")
            if os.path.exists(cand):
                with open(cand, encoding="utf-8", errors="replace") as f:
                    return f.read()
        except Exception:
            pass
    antiword = shutil.which("antiword")
    if antiword:
        try:
            r = subprocess.run([antiword, "-w", "0", path], capture_output=True, timeout=120)
            if r.returncode == 0 and r.stdout:
                return r.stdout.decode("utf-8", "replace")
        except Exception:
            pass
    return None


def _doc_builtin(path: str) -> str:
    """Last-resort pure-Python text scrape of an OLE .doc (imperfect but usable)."""
    with open(path, "rb") as f:
        data = f.read()
    # Heuristic: word text is mostly cp1252 or utf-16le runs
    runs = re.findall(rb"(?:[\x20-\x7e\r\x09][\x00]){4,}", data)  # utf-16le ascii runs
    if runs and sum(len(r) for r in runs) > len(data) // 20:
        text = b"".join(runs).decode("utf-16-le", "replace")
    else:
        runs = re.findall(rb"[\x20-\x7e\r\x09]{4,}", data)
        text = b"\n".join(runs).decode("cp1252", "replace")
    return text.replace("\r", "\n")


def to_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        with open(path, encoding="utf-8", errors="replace") as f:
            return f.read()
    if ext == ".docx":
        return _docx_text(path)
    if ext == ".doc":
        with tempfile.TemporaryDirectory() as td:
            t = _doc_via_tool(path, td)
        return t if t is not None else _doc_builtin(path)
    raise ValueError(f"unsupported extension: {path}")


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("input")
    ap.add_argument("-o", "--output")
    a = ap.parse_args()
    text = to_text(a.input)
    out = a.output or os.path.splitext(a.input)[0] + ".txt"
    with open(out, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"{a.input} -> {out} ({len(text)} chars)")


if __name__ == "__main__":
    sys.exit(main())
